import docx

# REQUIRED CLAUSES
REQUIRED_CLAUSES = {
    "data privacy": ["data privacy", "data protection", "gdpr"],
    "termination": ["termination", "cancel", "end of contract"],
    "governing law": ["governing law", "jurisdiction", "legal authority"],
    "payment terms": ["payment terms", "fees", "payment schedule"]
}


# ✅ Read DOCX content
def read_docx(filepath):
    doc = docx.Document(filepath)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.lower()


# ✅ Find missing clauses
def check_compliance(text):
    missing = []

    for clause, keywords in REQUIRED_CLAUSES.items():
        if not any(keyword in text for keyword in keywords):
            missing.append(clause)

    return missing


# ✅ Modify TXT files (optional)
def modify_txt(filepath, missing_clauses):
    with open(filepath, "a", encoding="utf-8") as f:
        for clause in missing_clauses:
            f.write(f"\n\n[ADDED] {clause.upper()} CLAUSE placeholder added by AI.\n")


# ✅ FIXED ✅ Modify DOCX files with 3 arguments
def modify_docx(input_path, output_path, missing_clauses):
    doc = docx.Document(input_path)

    for clause in missing_clauses:
        doc.add_heading(f"{clause.title()} Clause Added", level=2)
        doc.add_paragraph(
            f"This clause is automatically added because the original contract "
            f"was missing the '{clause}' requirement."
        )

    doc.save(output_path)
    return output_path
