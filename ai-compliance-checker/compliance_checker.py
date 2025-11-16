import os
import docx
from rapidfuzz import fuzz
from email_smtp import send_email
import gspread
from google.oauth2.service_account import Credentials
from dotenv import load_dotenv
from gspread_formatting import *

# -------------------------------
# ‚úÖ Load environment variables
# -------------------------------
load_dotenv()

GOOGLE_SHEET_ID = os.getenv("GOOGLE_SHEET_ID")
SERVICE_ACCOUNT_FILE = os.getenv("GOOGLE_SERVICE_CREDS")

if not SERVICE_ACCOUNT_FILE:
    raise ValueError("SERVICE_ACCOUNT_FILE not set. Check your .env file.")

SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]

# -------------------------------
# ‚úÖ Google Sheets Setup
# -------------------------------
creds = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
gs_client = gspread.authorize(creds)
sheet = gs_client.open_by_key(GOOGLE_SHEET_ID).sheet1

# -------------------------------
# ‚úÖ Clause Keywords
# -------------------------------
CLAUSE_KEYWORDS = {
    "Data Breach Notification": [
        "breach", "notify", "data breach", "incident report", "security breach", "breach notification"
    ],
    "Confidentiality": [
        "confidential", "non-disclosure", "NDA", "confidentiality", "proprietary information", "secret information"
    ],
    "Data Privacy Protection Right": [
        "right to access", "right to be forgotten", "data subject rights", "rectify", "erase data",
        "restrict processing", "data portability"
    ],
    "Business Associate Agreement": [
        "business associate agreement", "BAA", "protected health information", "PHI", "associate agreement"
    ],
    "Data Processing Agreement": [
        "data processing agreement", "processor", "controller", "processing purposes", "subprocessors",
        "processing activities"
    ],
}

CONTRACT_DIR = "contracts"

# -------------------------------
# ‚úÖ Extract text from .docx
# -------------------------------
def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts).lower()

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

# -------------------------------
# ‚úÖ Detect Clause Presence
# -------------------------------
def detect_present_clauses(text):
    found = set()
    for clause, keywords in CLAUSE_KEYWORDS.items():
        for kw in keywords:
            if fuzz.partial_ratio(kw.lower(), text) > 70:
                found.add(clause)
                break
    return list(found)

def find_missing_clauses(text):
    present = detect_present_clauses(text)
    return [clause for clause in CLAUSE_KEYWORDS if clause not in present]

# -------------------------------
# ‚úÖ Write Results to Google Sheets with coloring
# -------------------------------
def log_to_sheet(filename, present, missing):
    # Append row
    new_row_index = len(sheet.get_all_values()) + 1
    sheet.append_row([
        filename,
        ", ".join(present) if present else "None",
        ", ".join(missing) if missing else "None"
    ])

    # Apply formatting
    if missing:
        # Red background for issues
        fmt = CellFormat(backgroundColor=color(1, 0.6, 0.6))  # light red
    else:
        # Green background if all clauses present
        fmt = CellFormat(backgroundColor=color(0.6, 1, 0.6))  # light green

    # Apply formatting to the new row
    format_cell_range(sheet, f"A{new_row_index}:C{new_row_index}", fmt)

# -------------------------------
# ‚úÖ Main Compliance Checker
# -------------------------------
def check_compliance():
    # Load email settings
    recipients = os.getenv("EMAIL_TO").split(",")
    smtp_server = os.getenv("EMAIL_SMTP_HOST")
    smtp_port = int(os.getenv("EMAIL_SMTP_PORT"))
    smtp_user = os.getenv("EMAIL_FROM")
    smtp_password = os.getenv("EMAIL_PASSWORD")

    for file_name in os.listdir(CONTRACT_DIR):
        if not file_name.endswith(".docx"):
            continue

        path = os.path.join(CONTRACT_DIR, file_name)
        print(f"\nüîç Checking: {file_name}")

        text = extract_text_from_docx(path)
        present = detect_present_clauses(text)
        missing = find_missing_clauses(text)

        print(f"üìÑ Extracted Preview: {text[:150]}...\n")

        # ‚úÖ Log into Google Sheets with color
        log_to_sheet(file_name, present, missing)

        # ‚úÖ Email only if missing clauses
        if missing:
            print("‚ö†Ô∏è Missing Clauses:")
            for m in missing:
                print(f"- {m}")

            email_msg = f"Missing clauses in {file_name}:\n" + \
                        "\n".join([f"- {m}" for m in missing])

            send_email(
                subject="Contract Compliance Issues",
                body=email_msg,
                recipients=recipients,
                smtp_server=smtp_server,
                smtp_port=smtp_port,
                smtp_user=smtp_user,
                smtp_password=smtp_password
            )

        else:
            print("‚úÖ All clauses present.")

if __name__ == "__main__":
    check_compliance()
